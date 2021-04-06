from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

#添加各人员社保路径至列表中
social_basic_path = 'e:\\TestCode\\广州社保'
social_name_folder = os.listdir(social_basic_path)
social_paths_list = []
for i in social_name_folder:
    social_name_folder_path = social_basic_path + '\\' + i
    social_path = social_name_folder_path + '\\' + os.listdir(social_name_folder_path)[0]
    social_paths_list.append(social_path)

#打印人员社保路径
#for i in social_paths:
#    print(i)

document = Document('测试.docx')

#生成所需人员名字及证件列表
names_string = input('输入需要添加人员的名字（|作为分割线）：')
certificates_string = input('输入需要添加的证件（|作为分割线）：')
names = names_string.split('|')
certificates = certificates_string.split('|')

#打印所需人员名字及证件
#for name in names:
#    print(name)
#for certificate in certificates:
#    print(certificate)

#存放人员证件文件夹路径
basic_path = 'e:\\TestCode\\中量工程咨询有限公司'

for name in names:
    document.add_heading('（' + str(int(names.index(name))+1) + '）' + name)  #添加标题
    name_path = basic_path + '\\' + name      #具体人员路径
    name_certificates_list = os.listdir(name_path)    #具体人员证件列表
    for certificate in certificates:

        if certificate in '社保证明':       #判断证件是否为社保
            no_certificate = True           #人员社保是否存在的标记
            for social_path in social_paths_list:     #遍历社保列表，查看所需人员是否在社保列表中
                if name in social_path:
                    document.add_picture(social_path,Inches(6.299))
                    print(social_path)
                    no_certificate = False
                    continue
            if no_certificate:
                print(name + certificate + '不存在')
            continue
                
        no_certificate = True
        for name_certificate in name_certificates_list:          #遍历具体人员的具体证件
            if certificate in name_certificate:             #判断具体人员所需证件是否存在
                name_certificate_path = name_path + '\\' + name_certificate   #具体人员具体证件的路径
                if '身份证' in name_certificate_path:       #判断证件是否是身份证
                    paragraph = document.add_paragraph()             
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run("")
                    run.add_picture(name_certificate_path,Inches(3.4))           #添加身份证至文档中
                    #print(name + name_certificate + '已添加')
                    print(name_certificate_path)
                else:
                    document.add_picture(name_certificate_path,Inches(6.299))         #添加证件至文档中
                    #print(name + name_certificate + '已添加')
                    print(name_certificate_path)
                no_certificate = False
        if no_certificate:
            print(name + certificate + '不存在')

document.save('测试3.docx')
print('文件保存成功')
                
                
                
                   
    
        
    
